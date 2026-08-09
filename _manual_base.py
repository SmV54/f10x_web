# -*- coding: utf-8 -*-
"""Ferramentas de montagem dos manuais em Word.

Estilo, cabecalho de pagina e os blocos de texto ficam aqui, para os dois
manuais (o interno, de regras, e o do usuario) sairem com a mesma cara e
mudarem juntos. Cada script de manual so cuida do CONTEUDO.

Uso:
    from _manual_base import *
    iniciar("Folha10 Simples — Manual do Usuario", "versao 2026...")
    h1("Etapa 1"); p("texto"); tabela([...], [...])
    salvar(r"C:\\...\\arquivo.docx")
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL     = RGBColor(0x1D, 0x4E, 0xD8)
CINZA    = RGBColor(0x47, 0x55, 0x69)
VERMELHO = RGBColor(0xB9, 0x1C, 0x1C)
VERDE    = RGBColor(0x15, 0x80, 0x3D)

_doc = None


def documento():
    """O objeto Document, para o caso raro de precisar mexer direto."""
    return _doc


# ------------------------------------------------------------ montagem
def _borda_inferior(par, cor="C7D2E4"):
    pPr = par._p.get_or_add_pPr()
    bordas = OxmlElement("w:pBdr")
    linha = OxmlElement("w:bottom")
    linha.set(qn("w:val"), "single")
    linha.set(qn("w:sz"), "4")        # 4 oitavos de ponto = fio de 0,5pt
    linha.set(qn("w:space"), "4")
    linha.set(qn("w:color"), cor)
    bordas.append(linha)
    pPr.append(bordas)


def _sombrear(celula, cor_hex):
    tc = celula._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), cor_hex)
    tc.append(shd)


def iniciar(titulo_cabecalho, direita_cabecalho):
    """Cria o documento, aplica os estilos e monta o cabecalho das paginas."""
    global _doc
    _doc = Document()

    normal = _doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for nome, tam, cor in (("Heading 1", 17, AZUL),
                           ("Heading 2", 13, AZUL),
                           ("Heading 3", 11.5, CINZA)):
        st = _doc.styles[nome]
        st.font.name = "Calibri"
        st.font.size = Pt(tam)
        st.font.color.rgb = cor
        st.font.bold = True
        st.paragraph_format.space_before = Pt(14 if nome == "Heading 1" else 10)
        st.paragraph_format.space_after = Pt(6)

    sec = _doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    sec.left_margin = sec.right_margin = Cm(2.2)

    # Cabecalho: titulo a esquerda e versao a direita, com um tab alinhado a
    # direita na largura util. A capa sai limpa (primeira pagina diferente).
    sec.different_first_page_header_footer = True
    par = sec.header.paragraphs[0]
    par.text = ""
    par.paragraph_format.space_after = Pt(4)
    par.paragraph_format.tab_stops.add_tab_stop(
        sec.page_width - sec.left_margin - sec.right_margin, WD_TAB_ALIGNMENT.RIGHT)
    for txt in (titulo_cabecalho, "\t" + direita_cabecalho):
        r = par.add_run(txt)
        r.font.size = Pt(8.5)
        r.font.color.rgb = CINZA
    _borda_inferior(par)
    return _doc


def salvar(caminho):
    _doc.save(caminho)
    return caminho


# ------------------------------------------------------------- blocos
def h1(txt, quebra=True):
    if quebra:
        _doc.add_page_break()
    _doc.add_heading(txt, level=1)


def h2(txt):
    _doc.add_heading(txt, level=2)


def h3(txt):
    _doc.add_heading(txt, level=3)


def p(txt="", italico=False, cor=None, negrito=False):
    par = _doc.add_paragraph()
    run = par.add_run(txt)
    run.italic = italico
    run.bold = negrito
    if cor is not None:
        run.font.color.rgb = cor
    return par


def regra(titulo, texto):
    """Nome da regra em negrito, explicacao em seguida, no mesmo paragrafo."""
    par = _doc.add_paragraph()
    par.add_run(titulo + " — ").bold = True
    par.add_run(texto)
    return par


def item(txt, nivel=0):
    par = _doc.add_paragraph(txt, style="List Bullet")
    par.paragraph_format.left_indent = Cm(0.7 + 0.6 * nivel)
    par.paragraph_format.space_after = Pt(2)
    return par


def passo(numero, titulo, texto=""):
    """Passo numerado de um roteiro: numero e titulo em negrito, texto depois."""
    par = _doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(0.7)
    par.paragraph_format.space_after = Pt(4)
    r = par.add_run(f"{numero}. ")
    r.bold = True
    r.font.color.rgb = AZUL
    par.add_run(titulo).bold = True
    if texto:
        par.add_run("  " + texto)
    return par


def _faixa(rotulo, txt, cor, indent=0.4):
    par = _doc.add_paragraph()
    r = par.add_run(rotulo + "  ")
    r.bold = True
    r.font.color.rgb = cor
    par.add_run(txt)
    par.paragraph_format.left_indent = Cm(indent)
    return par


def atencao(txt):
    return _faixa("Atenção", txt, VERMELHO)


def dica(txt):
    return _faixa("Dica", txt, VERDE)


def tabela(cabecalho, linhas, larguras=None):
    t = _doc.add_table(rows=1, cols=len(cabecalho))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, c in enumerate(cabecalho):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(c)
        run.bold = True
        run.font.size = Pt(9.5)
        _sombrear(hdr[i], "E8EEF9")
    for ln in linhas:
        cells = t.add_row().cells
        for i, v in enumerate(ln):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(9.5)
    if larguras:
        for row in t.rows:
            for i, w in enumerate(larguras):
                row.cells[i].width = Cm(w)
    _doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def caminho(txt):
    """Linha de navegacao: onde a funcao fica no menu."""
    par = _doc.add_paragraph()
    r = par.add_run(txt)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = CINZA
    par.paragraph_format.space_after = Pt(8)
    return par


def capa(titulo, subtitulo, linha3=""):
    for txt, tam, cor, ital, neg in (
            (titulo, 30, AZUL, False, True),
            (subtitulo, 17, CINZA, False, False),
            (linha3, 11, CINZA, True, False)):
        if not txt:
            continue
        par = _doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = par.add_run(txt)
        r.font.size = Pt(tam)
        r.font.color.rgb = cor
        r.italic = ital
        r.bold = neg
    _doc.add_paragraph()
